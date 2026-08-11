#!/usr/bin/env python
"""Build the canonical final manuscript as DOCX and PDF.

The Markdown manuscript is the only maintained content source. This script
provides a small, deterministic renderer for the repository's Markdown subset
and validates figures, boundaries, provenance, and generated outputs.
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from html import escape as html_escape
import re
from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table as PdfTable,
    TableStyle,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = REPO_ROOT / "docs" / "report" / "final_report.md"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "dist"
DOCX_NAME = "Drosophila_PD_FlyGym_Final_Report.docx"
PDF_NAME = "Drosophila_PD_FlyGym_Final_Report.pdf"
SOURCE_COMMIT = "004488cf7fd5e980137a209d360b977716865e1a"
SOURCE_COMMIT_SHORT = "004488c"

EXPECTED_FIGURES = (
    "results/analysis/figures/e1_parameter_response.png",
    "results/analysis/figures/e2_condition_comparison.png",
    "results/analysis/figures/e3_paired_seed_robustness.png",
    "results/analysis/figures/e5_computational_reversibility.png",
)
REQUIRED_HEADINGS = (
    "# Reproducible Computational Drosophila Locomotion Phenotype Framework",
    "## Abstract",
    "## 1. Introduction",
    "## 2. Methods",
    "## 3. Results",
    "## 4. Discussion",
    "## 5. Limitations",
    "## 6. Reproducibility",
    "## 7. Conclusion",
    "## References",
    "## Appendix A. Frozen evidence chain",
    "## Appendix B. Reproduction commands",
    "## Appendix C. Evidence traceability",
)
BOUNDARY_TERMS = (
    "PARTIAL_PHENOTYPE_CONCORDANCE",
    "do not establish a biologically validated Parkinson's disease model",
    "not biological rescue",
    "No direct numerical calibration was used",
    "not biological recovery measurements",
    "No statistical-significance claim was introduced",
)


@dataclass(frozen=True)
class Block:
    kind: str
    value: object


def _rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _split_table_row(line: str) -> list[str]:
    parts = line.strip().split("|")
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [part.strip() for part in parts]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_markdown(path: Path) -> tuple[str, list[Block]]:
    """Parse the small Markdown subset used by the canonical manuscript."""

    lines = path.read_text(encoding="utf-8").splitlines()
    title = ""
    blocks: list[Block] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            text = " ".join(part.strip() for part in paragraph_lines).strip()
            if text:
                blocks.append(Block("paragraph", text))
            paragraph_lines.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            flush_paragraph()
            index += 1
            continue

        fence = re.match(r"^\s*```(?:[^`]*)$", line)
        if fence:
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(Block("code", "\n".join(code_lines)))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            heading_text = heading.group(2)
            if level == 1 and not title:
                title = heading_text
            blocks.append(Block("heading", (level, heading_text)))
            index += 1
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if image:
            flush_paragraph()
            blocks.append(Block("image", (image.group(1), image.group(2))))
            index += 1
            continue

        if line.lstrip().startswith("|"):
            flush_paragraph()
            table_lines = [line]
            index += 1
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [
                _split_table_row(table_line)
                for table_line in table_lines
                if not _is_table_separator(table_line)
            ]
            if rows:
                blocks.append(Block("table", rows))
            continue

        list_match = re.match(r"^\s*(-|\d+\.)\s+(.+)$", line)
        if list_match:
            flush_paragraph()
            ordered = list_match.group(1).endswith(".")
            items = [list_match.group(2).strip()]
            index += 1
            while index < len(lines):
                next_match = re.match(r"^\s*(-|\d+\.)\s+(.+)$", lines[index])
                if next_match:
                    if next_match.group(1).endswith(".") == ordered:
                        items.append(next_match.group(2).strip())
                        index += 1
                        continue
                    break
                if lines[index].strip() and not lines[index].lstrip().startswith(("#", "|", "!", "```")):
                    items[-1] = f"{items[-1]} {lines[index].strip()}"
                    index += 1
                    continue
                break
            list_kind = "ordered_list" if ordered else "bullet_list"
            if blocks and blocks[-1].kind == list_kind:
                blocks[-1] = Block(list_kind, [*blocks[-1].value, *items])
            else:
                blocks.append(Block(list_kind, items))
            continue

        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    if not title:
        raise ValueError("Canonical manuscript has no H1 title")
    return title, blocks


def _inline_tokens(text: str) -> Iterable[tuple[str, str]]:
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))"
    )
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            yield "plain", text[cursor : match.start()]
        token = match.group(0)
        if token.startswith("**"):
            yield "bold", token[2:-2]
        elif token.startswith("`"):
            yield "code", token[1:-1]
        elif token.startswith("*"):
            yield "italic", token[1:-1]
        else:
            yield "link", re.match(r"\[([^\]]+)\]", token).group(1)  # type: ignore[union-attr]
        cursor = match.end()
    if cursor < len(text):
        yield "plain", text[cursor:]


def _plain_inline(text: str) -> str:
    return "".join(value for _, value in _inline_tokens(text))


def _reportlab_inline(text: str) -> str:
    rendered: list[str] = []
    for kind, value in _inline_tokens(text):
        safe = html_escape(value)
        if kind == "bold":
            rendered.append(f"<b>{safe}</b>")
        elif kind == "italic":
            rendered.append(f"<i>{safe}</i>")
        elif kind == "code":
            rendered.append(f"<font name=Courier>{safe}</font>")
        else:
            rendered.append(safe)
    return "".join(rendered)


def _set_run_font(run, name: str, size: float, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.bold = bold
    run.italic = italic


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _add_inline_docx(paragraph, text: str, size: float = 11, color: str = "000000") -> None:
    for kind, value in _inline_tokens(text):
        run = paragraph.add_run(value)
        _set_run_font(run, "Calibri", size, color, bold=kind == "bold", italic=kind == "italic")
        if kind == "code":
            _set_run_font(run, "Consolas", size, color)


def _table_widths(column_count: int) -> list[int]:
    options = {
        3: [2200, 4500, 2660],
        4: [2100, 2200, 2250, 2810],
        6: [1800, 1300, 1300, 1300, 1300, 2360],
        8: [1100, 1150, 1150, 1150, 1150, 1150, 1300, 1260],
    }
    if column_count in options:
        return options[column_count]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def _resolve_manuscript_asset(relative_path: str) -> Path:
    return (DEFAULT_SOURCE.parent / relative_path).resolve()


def _configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb("000000")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, size, italic, color in (
        ("Report Caption", 9.5, True, "555555"),
        ("Report Table Text", 8.5, False, "000000"),
        ("Report Table Caption", 9.5, False, "1F4D78"),
        ("Report Code", 8.5, False, "333333"),
    ):
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri" if name != "Report Code" else "Consolas"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.italic = italic
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.space_after = Pt(4 if "Table" in name else 8)
        style.paragraph_format.line_spacing = 1.0 if "Table" in name or name == "Report Code" else 1.1


def _add_title_page(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(115)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(title)
    _set_run_font(run, "Calibri", 27, "0B2545", bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(32)
    run = paragraph.add_run("Milestone F | Final Submission Package")
    _set_run_font(run, "Calibri", 14, "2E74B5", bold=True)

    for label, value in (
        ("Canonical manuscript source", "docs/report/final_report.md"),
        ("Source commit", SOURCE_COMMIT),
        ("Scope", "Computational and phenomenological model only"),
        ("Evidence status", "E6 frozen evidence synthesis"),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(6)
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, "Calibri", 10.5, "555555", bold=True)
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, "Calibri", 10.5, "555555")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(42)
    run = paragraph.add_run(
        "PARTIAL_PHENOTYPE_CONCORDANCE remains qualitative; this package does "
        "not claim Parkinson's disease validation."
    )
    _set_run_font(run, "Calibri", 10, "7A5A00", italic=True)


def _add_contents_docx(document: Document, blocks: Sequence[Block]) -> None:
    paragraph = document.add_paragraph("Contents", style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    for block in blocks:
        if block.kind != "heading":
            continue
        level, text = block.value
        if level in (1, 2):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_inline_docx(paragraph, _plain_inline(text), 10.5, "1F4D78")


def _add_docx_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    widths = _table_widths(column_count)
    for row_index, row in enumerate(normalized):
        if row_index == 0:
            _set_repeat_header(table.rows[row_index])
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            _clear_paragraph(paragraph)
            paragraph.style = "Report Table Text"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_docx(paragraph, value, 8.5, "0B2545" if row_index == 0 else "000000")
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
            shading = cell._tc.get_or_add_tcPr().first_child_found_in("w:shd")
            if row_index == 0:
                if shading is None:
                    shading = OxmlElement("w:shd")
                    cell._tc.get_or_add_tcPr().append(shading)
                shading.set(qn("w:fill"), "F4F6F9")
    _set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_docx_block(document: Document, block: Block) -> None:
    if block.kind == "heading":
        level, text = block.value
        if level == 1 and text.startswith("Reproducible Computational"):
            return
        document.add_heading(_plain_inline(text), level=min(level, 3))
    elif block.kind == "paragraph":
        text = str(block.value)
        is_caption = text.startswith("*Figure ") and text.endswith("*")
        if is_caption:
            paragraph = document.add_paragraph(style="Report Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_docx(paragraph, text[1:-1], 9.5, "555555")
        else:
            paragraph = document.add_paragraph()
            if text.startswith("None of the present results establishes"):
                paragraph.paragraph_format.keep_with_next = True
            _add_inline_docx(paragraph, text)
    elif block.kind == "bullet_list" or block.kind == "ordered_list":
        style_name = "List Bullet" if block.kind == "bullet_list" else "List Number"
        for item in block.value:
            paragraph = document.add_paragraph(style=style_name)
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.167
            _add_inline_docx(paragraph, item)
    elif block.kind == "code":
        paragraph = document.add_paragraph(style="Report Code")
        _add_inline_docx(paragraph, block.value, 8.5, "333333")
    elif block.kind == "image":
        caption, relative_path = block.value
        image_path = _resolve_manuscript_asset(relative_path)
        with PILImage.open(image_path) as image:
            width, height = image.size
        target_width = 6.0
        target_height = target_width * height / width
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(target_width), height=Inches(target_height))
    elif block.kind == "table_caption":
        paragraph = document.add_paragraph(style="Report Table Caption")
        _add_inline_docx(paragraph, str(block.value), 9.5, "1F4D78")
    elif block.kind == "table":
        _add_docx_table(document, block.value)


def _set_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, "Calibri", 9, "777777")


def _configure_docx_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"Final report | source {SOURCE_COMMIT_SHORT}")
    _set_run_font(run, "Calibri", 8, "777777")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    label = footer.add_run("Page ")
    _set_run_font(label, "Calibri", 9, "777777")
    _set_page_field(footer)


def build_docx(title: str, blocks: Sequence[Block], output_path: Path) -> None:
    document = Document()
    _configure_docx_styles(document)
    _configure_docx_page(document)
    props = document.core_properties
    props.title = title
    props.subject = f"Final scientific report package; canonical source commit {SOURCE_COMMIT}"
    props.author = "Drosophila PD FlyGym repository"
    props.keywords = f"Drosophila, FlyGym, computational locomotion, source commit {SOURCE_COMMIT}"
    props.comments = "Built deterministically from docs/report/final_report.md."

    _add_title_page(document, title)
    document.add_page_break()
    _add_contents_docx(document, blocks)
    document.add_page_break()
    for block in blocks:
        _add_docx_block(document, block)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _pdf_styles():
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "ReportBody", parent=base["BodyText"], fontName="Helvetica", fontSize=10.5,
            leading=14, spaceAfter=8, alignment=TA_JUSTIFY, textColor=colors.HexColor("#000000"),
        ),
        "heading1": ParagraphStyle(
            "ReportH1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=16,
            leading=20, spaceBefore=18, spaceAfter=10, textColor=colors.HexColor("#2E74B5"),
            keepWithNext=True,
        ),
        "heading2": ParagraphStyle(
            "ReportH2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=13,
            leading=16, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#2E74B5"),
            keepWithNext=True,
        ),
        "heading3": ParagraphStyle(
            "ReportH3", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=12,
            leading=15, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#1F4D78"),
            keepWithNext=True,
        ),
        "caption": ParagraphStyle(
            "ReportCaption", parent=base["BodyText"], fontName="Helvetica-Oblique", fontSize=9,
            leading=11, spaceBefore=3, spaceAfter=8, alignment=TA_CENTER, textColor=colors.HexColor("#555555"),
        ),
        "table_caption": ParagraphStyle(
            "ReportTableCaption", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=9,
            leading=11, spaceBefore=5, spaceAfter=4, textColor=colors.HexColor("#1F4D78"),
        ),
        "table": ParagraphStyle(
            "ReportTable", parent=base["BodyText"], fontName="Helvetica", fontSize=7.2,
            leading=8.6, spaceAfter=0, alignment=TA_LEFT,
        ),
        "table_header": ParagraphStyle(
            "ReportTableHeader", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.2,
            leading=8.6, spaceAfter=0, alignment=TA_LEFT, textColor=colors.HexColor("#0B2545"),
        ),
        "code": ParagraphStyle(
            "ReportCode", parent=base["Code"], fontName="Courier", fontSize=8.3,
            leading=10, spaceAfter=8, leftIndent=12, rightIndent=12,
        ),
        "contents": ParagraphStyle(
            "ReportContents", parent=base["BodyText"], fontName="Helvetica", fontSize=10.5,
            leading=14, spaceAfter=4, textColor=colors.HexColor("#1F4D78"),
        ),
    }


def _pdf_table(rows: list[list[str]], styles) -> PdfTable:
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    widths_dxa = _table_widths(column_count)
    widths = [width / 1440 * inch for width in widths_dxa]
    data = []
    for row_index, row in enumerate(normalized):
        style = styles["table_header"] if row_index == 0 else styles["table"]
        data.append([Paragraph(_reportlab_inline(value), style) for value in row])
    table = PdfTable(data, colWidths=widths, repeatRows=1, hAlign="LEFT", splitByRow=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B7C0CC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def _add_pdf_title(story: list, title: str, styles) -> None:
    story.append(Spacer(1, 1.55 * inch))
    story.append(Paragraph(html_escape(title), ParagraphStyle(
        "CoverTitle", parent=styles["heading1"], fontName="Helvetica-Bold", fontSize=25,
        leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"),
        spaceBefore=0, spaceAfter=14,
    )))
    story.append(Paragraph("Milestone F | Final Submission Package", ParagraphStyle(
        "CoverSubtitle", parent=styles["body"], fontName="Helvetica-Bold", fontSize=14,
        leading=18, alignment=TA_CENTER, textColor=colors.HexColor("#2E74B5"), spaceAfter=28,
    )))
    for label, value in (
        ("Canonical manuscript source", "docs/report/final_report.md"),
        ("Source commit", SOURCE_COMMIT),
        ("Scope", "Computational and phenomenological model only"),
        ("Evidence status", "E6 frozen evidence synthesis"),
    ):
        story.append(Paragraph(
            f"<b>{html_escape(label)}:</b> {html_escape(value)}",
            ParagraphStyle("CoverMeta", parent=styles["body"], fontSize=10.5, leading=14,
                           alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=5),
        ))
    story.append(Spacer(1, 0.45 * inch))
    story.append(Paragraph(
        "<i>PARTIAL_PHENOTYPE_CONCORDANCE remains qualitative; this package does not claim Parkinson's disease validation.</i>",
        ParagraphStyle("CoverBoundary", parent=styles["body"], fontSize=10, leading=13,
                       alignment=TA_CENTER, textColor=colors.HexColor("#7A5A00"), spaceAfter=0),
    ))


def _add_pdf_contents(story: list, blocks: Sequence[Block], styles) -> None:
    story.append(Paragraph("Contents", styles["heading1"]))
    for block in blocks:
        if block.kind != "heading":
            continue
        level, text = block.value
        if level in (1, 2):
            style = styles["contents"]
            if level == 2:
                style = ParagraphStyle("ContentsIndented", parent=style, leftIndent=18)
            story.append(Paragraph(_reportlab_inline(_plain_inline(text)), style))


def _add_pdf_block(story: list, block: Block, styles) -> None:
    if block.kind == "heading":
        level, text = block.value
        if level == 1 and text.startswith("Reproducible Computational"):
            return
        story.append(Paragraph(_reportlab_inline(_plain_inline(text)), styles[f"heading{min(level, 3)}"]))
    elif block.kind == "paragraph":
        text = str(block.value)
        if text.startswith("*Figure ") and text.endswith("*"):
            story.append(Paragraph(_reportlab_inline(text[1:-1]), styles["caption"]))
        else:
            style = styles["body"]
            if text.startswith("None of the present results establishes"):
                style = ParagraphStyle("BoundaryTransition", parent=style, keepWithNext=True)
            story.append(Paragraph(_reportlab_inline(text), style))
    elif block.kind == "bullet_list" or block.kind == "ordered_list":
        for index, item in enumerate(block.value, start=1):
            marker = f"{index}. " if block.kind == "ordered_list" else "- "
            story.append(Paragraph(_reportlab_inline(marker + item), styles["body"]))
    elif block.kind == "code":
        story.append(Preformatted(str(block.value), styles["code"]))
    elif block.kind == "image":
        _, relative_path = block.value
        image_path = _resolve_manuscript_asset(relative_path)
        with PILImage.open(image_path) as image:
            width, height = image.size
        target_width = 6.0 * inch
        target_height = target_width * height / width
        image_flowable = PdfImage(str(image_path), width=target_width, height=target_height)
        image_flowable.hAlign = "CENTER"
        story.append(KeepTogether([image_flowable, Spacer(1, 5)]))
    elif block.kind == "table_caption":
        story.append(Paragraph(_reportlab_inline(str(block.value)), styles["table_caption"]))
    elif block.kind == "table":
        story.append(_pdf_table(block.value, styles))
        story.append(Spacer(1, 7))


def _draw_pdf_page(canvas, document) -> None:
    canvas.saveState()
    page_number = canvas.getPageNumber()
    if page_number > 1:
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawRightString(7.5 * inch, 10.48 * inch, f"Final report | source {SOURCE_COMMIT_SHORT}")
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#777777"))
    canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Page {page_number}")
    canvas.restoreState()


def build_pdf(title: str, blocks: Sequence[Block], output_path: Path) -> None:
    styles = _pdf_styles()
    story: list = []
    _add_pdf_title(story, title, styles)
    story.append(PageBreak())
    _add_pdf_contents(story, blocks, styles)
    story.append(PageBreak())
    for block in blocks:
        _add_pdf_block(story, block, styles)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path), pagesize=letter, leftMargin=1 * inch, rightMargin=1 * inch,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        title=title, author="Drosophila PD FlyGym repository",
        subject=f"Canonical source commit {SOURCE_COMMIT}; computational report package",
    )
    document.build(story, onFirstPage=_draw_pdf_page, onLaterPages=_draw_pdf_page)


def validate_source(source_path: Path) -> tuple[str, list[Block]]:
    if not source_path.exists():
        raise FileNotFoundError(source_path)
    source_text = source_path.read_text(encoding="utf-8")
    for heading in REQUIRED_HEADINGS:
        if heading not in source_text:
            raise ValueError(f"Missing required manuscript section: {heading}")
    for term in BOUNDARY_TERMS:
        if term not in source_text:
            raise ValueError(f"Missing scientific-boundary term: {term}")
    for relative_path in EXPECTED_FIGURES:
        if not (REPO_ROOT / relative_path).exists():
            raise FileNotFoundError(REPO_ROOT / relative_path)
    for relative_path in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", source_text):
        if relative_path.startswith(("http://", "https://")):
            continue
        asset = (source_path.parent / relative_path).resolve()
        if not asset.exists():
            raise FileNotFoundError(f"Missing local manuscript asset: {relative_path}")
    return parse_markdown(source_path)


def validate_outputs(docx_path: Path, pdf_path: Path) -> None:
    for path in (docx_path, pdf_path):
        if not path.exists() or path.stat().st_size == 0:
            raise ValueError(f"Missing or empty generated artifact: {path}")
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
        core_xml = archive.read("docProps/core.xml")
    if SOURCE_COMMIT.encode() not in document_xml and SOURCE_COMMIT.encode() not in core_xml:
        raise ValueError("DOCX does not contain canonical source commit metadata")
    pdf_bytes = pdf_path.read_bytes()
    if SOURCE_COMMIT.encode() not in pdf_bytes:
        raise ValueError("PDF does not contain canonical source commit metadata")


def build(source_path: Path = DEFAULT_SOURCE, output_dir: Path = DEFAULT_OUTPUT_DIR) -> tuple[Path, Path]:
    title, blocks = validate_source(source_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / DOCX_NAME
    pdf_path = output_dir / PDF_NAME
    build_docx(title, blocks, docx_path)
    build_pdf(title, blocks, pdf_path)
    validate_outputs(docx_path, pdf_path)
    return docx_path, pdf_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build the canonical final report as DOCX and PDF.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--validate-only", action="store_true", help="Validate source sections, boundaries, and local assets.")
    parser.add_argument("--check-outputs", action="store_true", help="Also require the target DOCX and PDF to exist and contain provenance metadata.")
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        validate_source(args.source)
        if args.validate_only:
            if args.check_outputs:
                validate_outputs(args.output_dir / DOCX_NAME, args.output_dir / PDF_NAME)
            print("Source validation: PASS")
            return 0
        docx_path, pdf_path = build(args.source, args.output_dir)
    except Exception as exc:
        print(f"Final report build failed: {type(exc).__name__}: {exc}")
        return 2
    print(f"Built DOCX: {docx_path}")
    print(f"Built PDF: {pdf_path}")
    print(f"Source commit: {SOURCE_COMMIT}")
    print("Output validation: PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
